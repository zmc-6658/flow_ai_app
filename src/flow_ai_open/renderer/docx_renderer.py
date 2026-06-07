from __future__ import annotations

from copy import deepcopy
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_BREAK
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Pt
from docx.text.paragraph import Paragraph
from lxml import etree

from flow_ai.core.ast_models import (
    DocumentAST,
    GeneratedAnchorNode,
    HeadingNode,
    OpaqueNode,
    OpaqueType,
    ParagraphNode,
    TextAlignment,
)
from flow_ai.core.preservation_models import AssetBlob, AssetStore, PreservationPlan
from flow_ai.core.profile_models import (
    FrontMatterProfile,
    PageBreakDecision,
    RenderProfiles,
)
from flow_ai.core.style_models import RenderPlan, StyleIntent


ALIGNMENT_MAP = {
    TextAlignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
    TextAlignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
    TextAlignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
    TextAlignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    TextAlignment.DISTRIBUTE: WD_ALIGN_PARAGRAPH.DISTRIBUTE,
}


class DocxRenderer:

    def __init__(
        self,
        *,
        preservation_plan: PreservationPlan | None = None,
        asset_store: AssetStore | None = None,
        template_path: str | Path | None = None,
        profiles: RenderProfiles | None = None,
    ) -> None:
        self.preservation_plan = preservation_plan or PreservationPlan()
        self.asset_store = asset_store or AssetStore()
        self.template_path = Path(template_path) if template_path is not None else None
        self.profiles = profiles or RenderProfiles.fallback()
        self._section_breaks_by_source_index: dict[int, list[str]] = {}
        self._inserted_section_break_indexes: set[int] = set()
        self._page_breaks_by_source_index: dict[int, int] = {}
        self._inserted_page_break_indexes: set[int] = set()
        self._toc_heading_paragraph_xml: str | None = None
        self._toc_body_paragraph_xml: list[str] = []
        # 整段回注的图片/表格已自带分页和分节，不能再次回放。
        self._embedded_break_source_indices: set[int] = set()
        self._mixed_source_indices: set[int] = set()

    def render(self, ast: DocumentAST, plan: RenderPlan, output_path: str) -> None:

        doc = self._new_document()
        self._mixed_source_indices = self._collect_mixed_source_indices(ast)
        self._ensure_toc_styles(doc)
        for node in ast.blocks:
            if getattr(node, "suppress_render", False) or self._is_plan_suppressed(
                plan, node
            ):
                self._append_due_page_breaks(doc, node)
                self._append_due_section_breaks(doc, node)
                continue
            intent = plan.node_styles.get(node.id) or self._fallback_intent_for(node)

            paragraph = self._render_node(doc, node, intent)
            if paragraph is not None:
                self._apply_style_intent(
                    doc,
                    paragraph,
                    intent,
                    node,
                    skip_generated_text=self._is_restored_image_node(node),
                )
            self._append_due_page_breaks(doc, node)
            self._append_due_section_breaks(doc, node)

        self._force_update_fields(doc)
        output = Path(output_path)
        output.parent.mkdir(parents=True, exist_ok=True)
        doc.save(output)

    def _new_document(self) -> Document:
        self._embedded_break_source_indices.clear()
        if self.template_path is None:
            return Document()

        doc = Document(str(self.template_path))
        self._section_breaks_by_source_index = self._collect_template_section_breaks(doc)
        self._page_breaks_by_source_index = self._collect_template_page_breaks(doc)
        self._toc_heading_paragraph_xml = self._collect_template_toc_heading(doc)
        self._toc_body_paragraph_xml = self._collect_template_toc_body(doc)
        self._inserted_section_break_indexes = set()
        self._inserted_page_break_indexes = set()
        self._clear_body_content(doc)
        return doc

    def _clear_body_content(self, doc: Document) -> None:
        if hasattr(doc._body, "clear_content"):
            doc._body.clear_content()
            return

        body = doc._body._element
        for child in list(body):
            if child.tag != qn("w:sectPr"):
                body.remove(child)

    def _collect_template_section_breaks(self, doc: Document) -> dict[int, list[str]]:
        section_breaks: dict[int, list[str]] = {}
        source_index = 0
        for child in doc._body._element.iterchildren():
            if child.tag == qn("w:sectPr"):
                continue
            if child.tag in (qn("w:p"), qn("w:tbl")):
                section_properties = child.find(qn("w:pPr"))
                section_properties = (
                    section_properties.find(qn("w:sectPr"))
                    if section_properties is not None
                    else None
                )
                if section_properties is not None:
                    section_breaks.setdefault(source_index, []).append(
                        etree.tostring(section_properties, encoding="unicode")
                    )
                source_index += 1
        return section_breaks

    def _collect_template_page_breaks(self, doc: Document) -> dict[int, int]:
        page_breaks: dict[int, int] = {}
        source_index = 0
        for child in doc._body._element.iterchildren():
            if child.tag == qn("w:sectPr"):
                continue
            if child.tag in (qn("w:p"), qn("w:tbl")):
                count = sum(
                    1
                    for item in child.iter(qn("w:br"))
                    if item.get(qn("w:type")) == "page"
                )
                if count:
                    page_breaks[source_index] = count
                source_index += 1
        return page_breaks

    def _collect_template_toc_heading(self, doc: Document) -> str | None:
        for paragraph in doc._body._element.xpath(".//w:p"):
            if self._is_toc_heading_xml(paragraph):
                return etree.tostring(paragraph, encoding="unicode")
        return None

    def _collect_template_toc_body(self, doc: Document) -> list[str]:
        toc_body: list[str] = []
        in_toc = False
        for paragraph in doc._body._element.xpath(".//w:p"):
            if not in_toc:
                in_toc = self._is_toc_heading_xml(paragraph)
                continue
            text = self._paragraph_xml_text(paragraph).strip()
            if not text:
                continue
            if not self._looks_like_toc_entry_xml(paragraph):
                break
            toc_body.append(etree.tostring(paragraph, encoding="unicode"))
        return toc_body

    def _is_toc_heading_xml(self, paragraph: object) -> bool:
        style_id = self._paragraph_xml_style_id(paragraph)
        if style_id in {"TOCHeading", "TOC10"}:
            return True
        return self._paragraph_xml_text(paragraph).replace(" ", "").strip() == "目录"

    def _looks_like_toc_entry_xml(self, paragraph: object) -> bool:
        if self._paragraph_xml_style_id(paragraph).upper().startswith("TOC"):
            return True
        text = self._paragraph_xml_text(paragraph).strip()
        return bool(text) and text[-1].isdigit()

    def _paragraph_xml_style_id(self, paragraph: object) -> str:
        style = paragraph.xpath("./w:pPr/w:pStyle/@w:val")
        return style[0] if style else ""

    def _paragraph_xml_text(self, paragraph: object) -> str:
        return "".join(text or "" for text in paragraph.xpath(".//w:t/text()"))

    def _append_due_section_breaks(self, doc: Document, node: object) -> None:
        source_index = getattr(node, "source_index", None)
        if source_index is None or source_index in self._inserted_section_break_indexes:
            return
        if source_index in self._embedded_break_source_indices:
            self._inserted_section_break_indexes.add(source_index)
            return

        for section_properties_xml in self._section_breaks_by_source_index.get(
            source_index, []
        ):
            paragraph = OxmlElement("w:p")
            paragraph_properties = OxmlElement("w:pPr")
            paragraph_properties.append(parse_xml(section_properties_xml))
            paragraph.append(paragraph_properties)
            self._append_body_element(doc, paragraph)
        self._inserted_section_break_indexes.add(source_index)

    def _append_due_page_breaks(self, doc: Document, node: object) -> None:
        source_index = getattr(node, "source_index", None)
        if source_index is None or source_index in self._inserted_page_break_indexes:
            return
        if source_index in self._embedded_break_source_indices:
            self._inserted_page_break_indexes.add(source_index)
            return
        decision = self.profiles.page_break_decisions.get(getattr(node, "id", ""))
        if decision == PageBreakDecision.DELETE:
            self._inserted_page_break_indexes.add(source_index)
            return

        for _ in range(self._page_breaks_by_source_index.get(source_index, 0)):
            paragraph = doc.add_paragraph("")
            paragraph.add_run().add_break(WD_BREAK.PAGE)
        self._inserted_page_break_indexes.add(source_index)

    def _mark_embedded_layout_from_template(self, node: object) -> None:
        """整段 OOXML 已回注，避免重复回放分页和分节。"""
        source_index = getattr(node, "source_index", None)
        if source_index is not None:
            self._embedded_break_source_indices.add(source_index)

    def _collect_mixed_source_indices(self, ast: DocumentAST) -> set[int]:
        counts: dict[int, int] = {}
        for block in ast.blocks:
            source_index = getattr(block, "source_index", None)
            if source_index is None:
                continue
            counts[source_index] = counts.get(source_index, 0) + 1
        return {source_index for source_index, count in counts.items() if count > 1}

    def _fallback_intent_for(self, node: object) -> StyleIntent:
        if isinstance(node, OpaqueNode):
            return StyleIntent(
                style_name="FlowOpaquePlaceholder",
                opaque_placeholder_text="[UNSUPPORTED OPAQUE: {opaque_type}]",
            )
        if isinstance(node, GeneratedAnchorNode):
            return StyleIntent(style_name="Normal")
        return StyleIntent(style_name="FlowBody")

    def _is_plan_suppressed(self, plan: RenderPlan, node: object) -> bool:
        trace = plan.rule_trace.get(getattr(node, "id", ""), [])
        return any(item.startswith("skip: suppress_render=True") for item in trace)

    def _render_node(
        self, doc: Document, node: object, intent: StyleIntent
    ) -> Paragraph | None:
        if isinstance(node, GeneratedAnchorNode):
            if node.anchor_type == "toc":
                self._append_toc_field(doc)
                return None
            return doc.add_paragraph("")

        if isinstance(node, OpaqueNode):
            is_preserved, paragraph = self._render_preserved_opaque(doc, node)
            if is_preserved:
                return paragraph
            return doc.add_paragraph("")

        if isinstance(node, HeadingNode | ParagraphNode):
            return self._render_text_node(
                doc,
                node,
                self._front_matter_profile_for(intent, node),
            )

        return None

    def _render_text_node(
        self,
        doc: Document,
        node: HeadingNode | ParagraphNode,
        front_profile: FrontMatterProfile | None = None,
    ) -> Paragraph:
        paragraph = doc.add_paragraph("")
        if not node.spans:
            return paragraph

        for text, features in self._iter_profiled_spans(node, front_profile):
            run = paragraph.add_run(text)
            self._apply_source_run_features(run, features)
        return paragraph

    def _iter_profiled_spans(
        self,
        node: HeadingNode | ParagraphNode,
        front_profile: FrontMatterProfile | None,
    ):
        if front_profile is None or not node.spans:
            for span in node.spans:
                yield span.text, span.features
            return

        prefix = front_profile.prefix_label
        first = node.spans[0]
        if first.text.startswith(prefix) and first.text != prefix:
            yield prefix, first.features
            yield first.text[len(prefix) :], first.features.model_copy(
                update={"bold": None, "font_size_pt": None}
            )
            for span in node.spans[1:]:
                yield span.text, span.features
            return

        for span in node.spans:
            yield span.text, span.features

    def _apply_source_run_features(self, run: object, features: object) -> None:
        if getattr(features, "bold", None) is not None:
            run.bold = features.bold
        if getattr(features, "italic", None) is not None:
            run.italic = features.italic
        if getattr(features, "underline", None) is not None:
            run.underline = features.underline
        if getattr(features, "font_family", None) is not None:
            run.font.name = features.font_family
            r_fonts = self._get_or_add_r_fonts(run)
            r_fonts.set(qn("w:ascii"), features.font_family)
            r_fonts.set(qn("w:hAnsi"), features.font_family)
            r_fonts.set(qn("w:eastAsia"), features.font_family)
        if getattr(features, "font_size_pt", None) is not None:
            run.font.size = Pt(features.font_size_pt)

    def _render_preserved_opaque(
        self, doc: Document, node: OpaqueNode
    ) -> tuple[bool, Paragraph | None]:
        target = self.preservation_plan.target_for(node.id)
        if target is None:
            return False, None

        asset = self.asset_store.get(target.asset_id)
        if asset is None:
            return False, None

        if node.opaque_type is OpaqueType.IMAGE:
            source_index = getattr(node, "source_index", None)
            if source_index in self._mixed_source_indices:
                paragraph = self._render_image_payload(doc, asset)
                return paragraph is not None, paragraph
            is_rendered, paragraph = self._render_image_asset(doc, asset)
            if is_rendered and asset.xml is not None and asset.source_relationship_id is not None:
                self._mark_embedded_layout_from_template(node)
            return is_rendered, paragraph
        if node.opaque_type is OpaqueType.TABLE:
            if asset.xml is not None:
                self._append_table_xml(doc, asset)
                self._mark_embedded_layout_from_template(node)
                return True, None
            return False, None
        return False, None

    def _render_image_asset(
        self, doc: Document, asset: AssetBlob
    ) -> tuple[bool, Paragraph | None]:
        if asset.payload is None:
            return False, None

        if asset.xml is not None and asset.source_relationship_id is not None:
            self._append_image_paragraph_xml(doc, asset)
            return True, None

        paragraph = self._render_image_payload(doc, asset)
        return paragraph is not None, paragraph

    def _render_image_payload(self, doc: Document, asset: AssetBlob) -> Paragraph | None:
        if asset.payload is None:
            return None
        paragraph = doc.add_paragraph("")
        paragraph.add_run().add_picture(BytesIO(asset.payload))
        return paragraph

    def _append_image_paragraph_xml(self, doc: Document, asset: AssetBlob) -> None:
        temp_paragraph = doc.add_paragraph("")
        temp_paragraph.add_run().add_picture(BytesIO(asset.payload or b""))
        new_relationship_id = self._image_relationship_id(temp_paragraph)
        temp_paragraph._element.getparent().remove(temp_paragraph._element)
        if new_relationship_id is None:
            return

        paragraph_element = parse_xml(asset.xml or "")
        self._replace_relationship_id(
            paragraph_element,
            old_relationship_id=asset.source_relationship_id or "",
            new_relationship_id=new_relationship_id,
        )
        self._append_body_element(doc, paragraph_element)

    def _append_table_xml(self, doc: Document, asset: AssetBlob) -> None:
        if asset.xml is None:
            return

        table_element = parse_xml(asset.xml)
        self._append_body_element(doc, table_element)

    def _append_body_element(self, doc: Document, element: object) -> None:
        body = doc._body._element
        section_properties = body.sectPr
        if section_properties is None:
            body.append(element)
            return
        body.insert(body.index(section_properties), element)

    def _image_relationship_id(self, paragraph: Paragraph) -> str | None:
        for blip in paragraph._element.iter(qn("a:blip")):
            relationship_id = blip.get(qn("r:embed")) or blip.get(qn("r:link"))
            if relationship_id:
                return relationship_id
        return None

    def _replace_relationship_id(
        self,
        element: object,
        *,
        old_relationship_id: str,
        new_relationship_id: str,
    ) -> None:
        relationship_attrs = (qn("r:embed"), qn("r:link"), qn("r:id"))
        for child in element.iter():
            for attr_name in relationship_attrs:
                if child.get(attr_name) == old_relationship_id:
                    child.set(attr_name, new_relationship_id)

    def _append_toc_field(self, doc: Document) -> None:
        toc_heading_xml = (
            self._toc_heading_paragraph_xml
            or """
                <w:p>
                  <w:pPr>
                    <w:pStyle w:val="TOCHeading"/>
                    <w:jc w:val="center"/>
                  </w:pPr>
                  <w:r>
                    <w:rPr>
                      <w:rFonts w:ascii="SimHei" w:hAnsi="SimHei" w:eastAsia="SimHei"/>
                      <w:sz w:val="36"/>
                    </w:rPr>
                    <w:t>目 录</w:t>
                  </w:r>
                </w:p>
                  <w:r>
                    <w:t xml:space="preserve"></w:t>
                  </w:r>"""
        )
        sdt_xml = f"""<w:sdt {nsdecls("w")}>
              <w:sdtPr>
                <w:docPartObj>
                  <w:docPartGallery w:val="Table of Contents"/>
                  <w:docPartUnique/>
                </w:docPartObj>
              </w:sdtPr>
              <w:sdtContent>
                {toc_heading_xml}
                <w:p>
                  <w:r>
                    <w:fldChar w:fldCharType="begin"/>
                  </w:r>
                  <w:r>
                    <w:instrText xml:space="preserve">TOC \\o "1-3" \\h \\z \\u</w:instrText>
                  </w:r>
                  <w:r>
                    <w:fldChar w:fldCharType="separate"/>
                  </w:r>
                </w:p>
                <w:p>
                  <w:r>
                    <w:fldChar w:fldCharType="end"/>
                  </w:r>
                </w:p>
              </w:sdtContent>
            </w:sdt>"""
        self._append_body_element(doc, parse_xml(sdt_xml))

    def _ensure_toc_styles(self, doc: Document) -> None:
        tab_position = self.profiles.toc_visual.right_tab_twips or self._toc_right_tab_position(doc)
        for level in range(1, 4):
            style_name = self.profiles.toc_visual.entry_styles.get(level, f"TOC {level}")
            style = self._get_or_create_paragraph_style(doc, style_name)
            self._set_toc_style_paragraph_xml(
                style,
                tab_position=tab_position,
                left_indent=(level - 1) * 240,
                has_dot_leader=self.profiles.toc_visual.has_dot_leader,
            )

    def _toc_right_tab_position(self, doc: Document) -> int:
        section = doc.sections[0]
        page_width = getattr(section.page_width, "twips", None)
        left_margin = getattr(section.left_margin, "twips", None)
        right_margin = getattr(section.right_margin, "twips", None)
        if page_width is None or left_margin is None or right_margin is None:
            return 9360
        return max(720, page_width - left_margin - right_margin)

    def _set_toc_style_paragraph_xml(
        self,
        style: object,
        *,
        tab_position: int,
        left_indent: int,
        has_dot_leader: bool,
    ) -> None:
        paragraph_properties = self._get_or_add_style_p_pr(style)

        existing_tabs_el = paragraph_properties.find(qn("w:tabs"))
        tabs = OxmlElement("w:tabs")
        if existing_tabs_el is not None:
            for tab_elem in existing_tabs_el.findall(qn("w:tab")):
                if tab_elem.get(qn("w:val")) != "right":
                    tabs.append(deepcopy(tab_elem))
            paragraph_properties.remove(existing_tabs_el)

        right_tab = OxmlElement("w:tab")
        right_tab.set(qn("w:val"), "right")
        if has_dot_leader:
            right_tab.set(qn("w:leader"), "dot")
        right_tab.set(qn("w:pos"), str(tab_position))
        tabs.append(right_tab)
        paragraph_properties.append(tabs)

        if paragraph_properties.find(qn("w:ind")) is None:
            indentation = OxmlElement("w:ind")
            indentation.set(qn("w:left"), str(left_indent))
            paragraph_properties.append(indentation)

    def _is_restored_image_node(self, node: object) -> bool:
        if not isinstance(node, OpaqueNode) or node.opaque_type is not OpaqueType.IMAGE:
            return False

        target = self.preservation_plan.target_for(node.id)
        if target is None:
            return False

        asset = self.asset_store.get(target.asset_id)
        return asset is not None and asset.payload is not None

    def _front_matter_profile_for(
        self,
        intent: StyleIntent,
        node: object,
    ) -> FrontMatterProfile | None:
        role = getattr(node, "semantic_role", None)
        if role in self.profiles.front_matter:
            return self.profiles.front_matter[role]

        for profile in self.profiles.front_matter.values():
            if intent.style_name == profile.content_style:
                return profile
        return None

    def _apply_style_intent(
        self,
        doc: Document,
        paragraph: Paragraph,
        intent: StyleIntent,
        node: object,
        *,
        skip_generated_text: bool = False,
    ) -> None:
        if not skip_generated_text:
            self._apply_generated_text(paragraph, intent, node)
        if intent.style_name is not None:
            self._ensure_paragraph_style(doc, intent)
            paragraph.style = intent.style_name

        paragraph_format = paragraph.paragraph_format

        if intent.page_break_before is True:
            paragraph_format.page_break_before = True
        if intent.keep_with_next is not None:
            paragraph_format.keep_with_next = intent.keep_with_next
        if intent.alignment is not None and intent.alignment in ALIGNMENT_MAP:
            paragraph.alignment = ALIGNMENT_MAP[intent.alignment]
        if intent.space_before_pt is not None:
            paragraph_format.space_before = Pt(intent.space_before_pt)
        if intent.space_after_pt is not None:
            paragraph_format.space_after = Pt(intent.space_after_pt)
        if intent.first_line_indent_pt is not None:
            paragraph_format.first_line_indent = Pt(intent.first_line_indent_pt)
        if intent.hanging_indent_pt is not None:
            paragraph_format.first_line_indent = Pt(-intent.hanging_indent_pt)
        if intent.left_indent_pt is not None:
            paragraph_format.left_indent = Pt(intent.left_indent_pt)
        if intent.right_indent_pt is not None:
            paragraph_format.right_indent = Pt(intent.right_indent_pt)
        if intent.line_spacing_multiple is not None:
            paragraph_format.line_spacing = intent.line_spacing_multiple
        elif intent.line_spacing_pt is not None:
            paragraph_format.line_spacing = Pt(intent.line_spacing_pt)
        elif intent.line_spacing is not None:
            paragraph_format.line_spacing = intent.line_spacing
        if intent.outline_level is not None:
            self._set_outline_level(paragraph, intent.outline_level)

        for run in paragraph.runs:
            ascii_font = intent.ascii_font or intent.font_name
            hansi_font = intent.hansi_font or intent.font_name
            east_asia_font = intent.east_asia_font or intent.font_name
            if intent.font_name is not None:
                run.font.name = intent.font_name
            r_fonts = self._get_or_add_r_fonts(run)
            if ascii_font is not None:
                r_fonts.set(qn("w:ascii"), ascii_font)
            if hansi_font is not None:
                r_fonts.set(qn("w:hAnsi"), hansi_font)
            if east_asia_font is not None:
                r_fonts.set(qn("w:eastAsia"), east_asia_font)
            if intent.font_size_pt is not None:
                run.font.size = Pt(intent.font_size_pt)
            if intent.bold is not None:
                run.bold = intent.bold
        self._apply_front_matter_label_format(
            paragraph,
            self._front_matter_profile_for(intent, node),
            intent,
        )

    def _apply_front_matter_label_format(
        self,
        paragraph: Paragraph,
        profile: FrontMatterProfile | None,
        intent: StyleIntent,
    ) -> None:
        if profile is None:
            return
        if not paragraph.runs:
            return

        label_run = paragraph.runs[0]
        label_run.bold = profile.label_font_bold
        if profile.label_font_size_pt is not None:
            label_run.font.size = Pt(profile.label_font_size_pt)
        if profile.label_font_name is not None:
            label_run.font.name = profile.label_font_name
        r_fonts = self._get_or_add_r_fonts(label_run)
        if profile.label_font_name is not None:
            r_fonts.set(qn("w:ascii"), profile.label_font_name)
            r_fonts.set(qn("w:hAnsi"), profile.label_font_name)
            r_fonts.set(qn("w:eastAsia"), profile.label_font_name)
        if profile.no_first_line_indent:
            paragraph.paragraph_format.first_line_indent = Pt(0)

    def _apply_generated_text(
        self, paragraph: Paragraph, intent: StyleIntent, node: object
    ) -> None:
        text = intent.generated_anchor_text or intent.opaque_placeholder_text
        if text is not None:
            paragraph.text = text.format(
                anchor_type=getattr(node, "anchor_type", ""),
                opaque_type=getattr(getattr(node, "opaque_type", ""), "value", ""),
                node_id=getattr(node, "id", ""),
            )

    def _ensure_paragraph_style(self, doc: Document, intent: StyleIntent) -> None:
        if intent.style_name is None:
            return

        styles = doc.styles
        style = self._get_or_create_paragraph_style(doc, intent.style_name)

        if intent.font_name is not None:
            style.font.name = intent.font_name
        if intent.font_size_pt is not None:
            style.font.size = Pt(intent.font_size_pt)
        if intent.bold is not None:
            style.font.bold = intent.bold
        if intent.outline_level is not None:
            self._set_style_outline_level(style, intent.outline_level)
        self._apply_style_r_fonts(style, intent)
        self._apply_style_paragraph_format(style, intent)

    def _get_or_create_paragraph_style(self, doc: Document, style_name: str) -> object:
        styles = doc.styles
        if style_name in styles:
            return styles[style_name]
        return styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)

    def _get_or_add_r_fonts(self, run: object) -> object:
        run_properties = run._element.get_or_add_rPr()
        r_fonts = run_properties.rFonts
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            run_properties.insert(0, r_fonts)
        return r_fonts

    def _set_outline_level(self, paragraph: Paragraph, level: int) -> None:
        paragraph_properties = paragraph._p.get_or_add_pPr()
        outline_level = paragraph_properties.find(qn("w:outlineLvl"))
        if outline_level is None:
            outline_level = OxmlElement("w:outlineLvl")
            paragraph_properties.append(outline_level)
        outline_level.set(qn("w:val"), str(level))

    def _set_style_outline_level(self, style: object, level: int) -> None:
        paragraph_properties = self._get_or_add_style_p_pr(style)
        outline_level = paragraph_properties.find(qn("w:outlineLvl"))
        if outline_level is None:
            outline_level = OxmlElement("w:outlineLvl")
            paragraph_properties.append(outline_level)
        outline_level.set(qn("w:val"), str(level))

    def _apply_style_r_fonts(self, style: object, intent: StyleIntent) -> None:
        style_element = style.element
        run_properties = style_element.find(qn("w:rPr"))
        if run_properties is None:
            run_properties = OxmlElement("w:rPr")
            style_element.append(run_properties)

        r_fonts = run_properties.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            run_properties.insert(0, r_fonts)

        ascii_font = intent.ascii_font or intent.font_name
        hansi_font = intent.hansi_font or intent.font_name
        east_asia_font = intent.east_asia_font or intent.font_name
        if ascii_font is not None:
            r_fonts.set(qn("w:ascii"), ascii_font)
        if hansi_font is not None:
            r_fonts.set(qn("w:hAnsi"), hansi_font)
        if east_asia_font is not None:
            r_fonts.set(qn("w:eastAsia"), east_asia_font)

    def _apply_style_paragraph_format(self, style: object, intent: StyleIntent) -> None:
        paragraph_format = style.paragraph_format
        if intent.keep_with_next is not None:
            paragraph_format.keep_with_next = intent.keep_with_next
        if intent.alignment is not None and intent.alignment in ALIGNMENT_MAP:
            paragraph_format.alignment = ALIGNMENT_MAP[intent.alignment]
        if intent.space_before_pt is not None:
            paragraph_format.space_before = Pt(intent.space_before_pt)
        if intent.space_after_pt is not None:
            paragraph_format.space_after = Pt(intent.space_after_pt)
        if intent.first_line_indent_pt is not None:
            paragraph_format.first_line_indent = Pt(intent.first_line_indent_pt)
        if intent.hanging_indent_pt is not None:
            paragraph_format.first_line_indent = Pt(-intent.hanging_indent_pt)
        if intent.left_indent_pt is not None:
            paragraph_format.left_indent = Pt(intent.left_indent_pt)
        if intent.right_indent_pt is not None:
            paragraph_format.right_indent = Pt(intent.right_indent_pt)
        if intent.line_spacing_multiple is not None:
            paragraph_format.line_spacing = intent.line_spacing_multiple
        elif intent.line_spacing_pt is not None:
            paragraph_format.line_spacing = Pt(intent.line_spacing_pt)
        elif intent.line_spacing is not None:
            paragraph_format.line_spacing = intent.line_spacing

    def _get_or_add_style_p_pr(self, style: object) -> object:
        paragraph_properties = style.element.find(qn("w:pPr"))
        if paragraph_properties is None:
            paragraph_properties = OxmlElement("w:pPr")
            run_properties = style.element.find(qn("w:rPr"))
            if run_properties is None:
                style.element.append(paragraph_properties)
            else:
                style.element.insert(style.element.index(run_properties), paragraph_properties)
        return paragraph_properties

    def _force_update_fields(self, doc: Document) -> None:
        settings = doc.settings.element
        update_fields = settings.find(qn("w:updateFields"))
        if update_fields is None:
            update_fields = OxmlElement("w:updateFields")
            settings.append(update_fields)
        update_fields.set(qn("w:val"), "true")
