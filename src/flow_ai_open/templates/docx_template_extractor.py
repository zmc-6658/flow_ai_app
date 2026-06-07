from __future__ import annotations

import re
import sys
from collections import Counter
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any

import yaml
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Length
from docx.text.paragraph import Paragraph

if __package__ in (None, ""):
    sys.path.append(str(Path(__file__).resolve().parents[2]))

from flow_ai.core.ast_models import SemanticRole
from flow_ai.core.profile_models import (
    AcknowledgmentProfile,
    FrontMatterProfile,
    ReferenceProfile,
    RenderProfiles,
    TocVisualProfile,
)
from flow_ai.core.style_models import RuleNode
from flow_ai_open.templates.style_resolver import ResolvedStyle, WordStyleResolver
from flow_ai_open.utils.typography_cn import (
    line_spacing_to_cn,
    pt_to_cn_chars,
    pt_to_cn_size,
)


ALIGNMENT_TO_VALUE = {
    WD_ALIGN_PARAGRAPH.LEFT: "left",
    WD_ALIGN_PARAGRAPH.CENTER: "center",
    WD_ALIGN_PARAGRAPH.RIGHT: "right",
    WD_ALIGN_PARAGRAPH.JUSTIFY: "justify",
    WD_ALIGN_PARAGRAPH.DISTRIBUTE: "distribute",
}

H1_RE = re.compile(
    r"^(第[一二三四五六七八九十百千万\d]+[章节篇]|[一二三四五六七八九十]+[、\s]+[\u4e00-\u9fffA-Za-z]|[1-9]\d*[、．]\s*\S+|[1-9]\d*\.\s+\S+)"
)
H2_RE = re.compile(r"^([1-9]\d*\.[1-9]\d*(?:\s+|\S)|（[一二三四五六七八九十]+）|\([一二三四五六七八九十]+\))")
REFERENCE_RE = re.compile(r"^参\s*考\s*文\s*献$")
REFERENCE_ITEM_RE = re.compile(r"^(\[\d+\]|［\d+］|\d+[.．、]\s*)")
EXCLUDED_TITLE_RE = re.compile(r"^(摘要|目录|参考文献|致谢|附录|关键词|ABSTRACT|Abstract)$", re.IGNORECASE)


@dataclass(frozen=True)
class SampledParagraph:
    role: str
    paragraph: Paragraph
    index: int
    text: str


@dataclass(frozen=True)
class RuleProfileEntry:

    id: str
    label: str
    description: str
    selector: dict[str, Any]
    intent: dict[str, Any]
    readable: list[str]
    sample_text: str
    sample_index: int
    confidence: float
    evidence: list[str] = field(default_factory=list)
    sources: dict[str, str] = field(default_factory=dict)
    style_chain: list[str] = field(default_factory=list)
    priority: int = 1

    def to_profile_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "label": self.label,
            "description": self.description,
            "selector": self.selector,
            "intent": self.intent,
            "readable": self.readable,
            "sample": {
                "index": self.sample_index,
                "text": self.sample_text,
            },
            "confidence": self.confidence,
            "evidence": self.evidence,
            "sources": self.sources,
            "style_chain": self.style_chain,
            "priority": self.priority,
        }

    def to_rule_dict(self) -> dict[str, Any]:
        return {
            "id": self.id,
            "selector": self.selector,
            "apply": self.intent,
            "priority": self.priority,
        }


def normalize_text(text: str) -> str:

    return re.sub(r"\s+", " ", text).strip()


def extract_physical_properties(
    paragraph: Paragraph,
    resolver: WordStyleResolver | None = None,
) -> dict[str, Any]:

    if resolver is not None:
        return resolver.resolve_paragraph(paragraph).intent

    properties: dict[str, Any] = {}
    paragraph_format = paragraph.paragraph_format
    style_format = paragraph.style.paragraph_format if paragraph.style is not None else None

    alignment = paragraph.alignment
    if alignment is None and style_format is not None:
        alignment = style_format.alignment
    if alignment in ALIGNMENT_TO_VALUE:
        properties["alignment"] = ALIGNMENT_TO_VALUE[alignment]

    _put_pt(properties, "space_before_pt", paragraph_format.space_before)
    _put_pt(properties, "space_after_pt", paragraph_format.space_after)
    _put_pt(properties, "first_line_indent_pt", paragraph_format.first_line_indent)
    _put_pt(properties, "left_indent_pt", paragraph_format.left_indent)
    _put_pt(properties, "right_indent_pt", paragraph_format.right_indent)

    if "space_before_pt" not in properties and style_format is not None:
        _put_pt(properties, "space_before_pt", style_format.space_before)
    if "space_after_pt" not in properties and style_format is not None:
        _put_pt(properties, "space_after_pt", style_format.space_after)
    if "first_line_indent_pt" not in properties and style_format is not None:
        _put_pt(properties, "first_line_indent_pt", style_format.first_line_indent)

    hanging_indent = properties.get("first_line_indent_pt")
    if hanging_indent is not None and hanging_indent < 0:
        properties.pop("first_line_indent_pt")
        properties["hanging_indent_pt"] = abs(hanging_indent)

    line_spacing = paragraph_format.line_spacing
    if line_spacing is None and style_format is not None:
        line_spacing = style_format.line_spacing
    _put_line_spacing(properties, line_spacing)

    page_break_before = paragraph_format.page_break_before
    if page_break_before is None and style_format is not None:
        page_break_before = style_format.page_break_before
    if page_break_before is not None:
        properties["page_break_before"] = bool(page_break_before)

    keep_with_next = paragraph_format.keep_with_next
    if keep_with_next is None and style_format is not None:
        keep_with_next = style_format.keep_with_next
    if keep_with_next is not None:
        properties["keep_with_next"] = bool(keep_with_next)

    outline_level = _extract_outline_level(paragraph)
    if outline_level is not None:
        properties["outline_level"] = outline_level

    font_properties = _extract_font_properties(paragraph)
    properties.update(font_properties)

    return _drop_none(properties)


def sniff_semantic_samples(document: Document) -> dict[str, SampledParagraph]:

    paragraphs = [
        (index, paragraph, normalize_text(paragraph.text))
        for index, paragraph in enumerate(document.paragraphs)
        if normalize_text(paragraph.text)
    ]
    samples: dict[str, SampledParagraph] = {}

    title = _find_title_cn(paragraphs)
    if title is not None:
        samples["title_cn"] = title

    for index, paragraph, text in paragraphs[:40]:
        role = _front_matter_sample_role(text, title is not None)
        if role is not None and role not in samples:
            samples[role] = SampledParagraph(role, paragraph, index, text)

    references_index: int | None = None
    h1_index: int | None = None
    for index, paragraph, text in paragraphs:
        if "h1" not in samples and _looks_like_h1(paragraph, text):
            samples["h1"] = SampledParagraph("h1", paragraph, index, text)
            h1_index = index
        if "references_heading" not in samples and REFERENCE_RE.match(text):
            samples["references_heading"] = SampledParagraph(
                "references_heading", paragraph, index, text
            )
            references_index = index

        if len({"h1", "references_heading"} - samples.keys()) == 0:
            break

    h2 = _find_heading_sample(paragraphs, h1_index, references_index, level=2)
    if h2 is not None:
        samples["h2"] = h2

    body = _find_body_sample(paragraphs, h1_index, references_index)
    if body is not None:
        samples["body"] = body

    if references_index is not None:
        for index, paragraph, text in paragraphs:
            if index <= references_index:
                continue
            if _looks_like_reference_item(text):
                samples["back_paragraph"] = SampledParagraph(
                    "back_paragraph", paragraph, index, text
                )
                break

    if "back_paragraph" not in samples:
        for index, paragraph, text in paragraphs:
            if _looks_like_reference_item(text):
                samples["back_paragraph"] = SampledParagraph(
                    "back_paragraph", paragraph, index, text
                )
                break

    return samples


def build_rule_profile(
    samples: dict[str, SampledParagraph],
    resolver: WordStyleResolver,
) -> list[RuleProfileEntry]:

    entries: list[RuleProfileEntry] = []
    for role, spec in _role_specs().items():
        sample = samples.get(role)
        if sample is None:
            continue

        resolved = resolver.resolve_paragraph(sample.paragraph)
        intent = dict(resolved.intent)
        intent["style_name"] = spec["style_name"]
        if "outline_level" in spec:
            intent["outline_level"] = spec["outline_level"]

        entries.append(
            RuleProfileEntry(
                id=spec["id"],
                label=spec["label"],
                description=spec["description"],
                selector=spec["selector"],
                intent=intent,
                readable=_intent_to_readable(intent),
                sample_text=sample.text,
                sample_index=sample.index,
                confidence=spec["confidence"],
                evidence=[
                    f"嗅探角色: {role}",
                    f"样本段落序号: {sample.index}",
                    f"样式继承链: {' -> '.join(resolved.style_chain) or '无'}",
                ],
                sources=resolved.sources,
                style_chain=resolved.style_chain,
                priority=spec["priority"],
            )
        )
    return entries


def build_rules_from_samples(
    samples: dict[str, SampledParagraph],
    resolver: WordStyleResolver | None = None,
) -> list[dict[str, Any]]:

    if resolver is None:
        first_sample = next(iter(samples.values()), None)
        if first_sample is None:
            return _validated_rule_dicts(_generated_support_rules())
        resolver = WordStyleResolver(first_sample.paragraph.part.document)

    entries = build_rule_profile(samples, resolver)
    rules = [entry.to_rule_dict() for entry in entries]
    rules.extend(_missing_front_matter_default_rules(rules))
    rules.extend(_generated_support_rules())
    return _validated_rule_dicts(rules)


def build_profile_document(
    samples: dict[str, SampledParagraph],
    resolver: WordStyleResolver,
    document: Document | None = None,
) -> dict[str, Any]:

    entries = build_rule_profile(samples, resolver)
    rules = [entry.to_rule_dict() for entry in entries]
    rules.extend(_missing_front_matter_default_rules(rules))
    rules.extend(_generated_support_rules())
    return {
        "profile": {
            "schema_version": "rule_profile.v1",
            "title": "Extracted thesis rule profile",
            "description": "机器从 Word 范文样式继承链提取的候选排版规则，供前端确权。",
            "rules": [entry.to_profile_dict() for entry in entries],
            "warnings": _profile_warnings(entries),
        },
        "profiles": _build_render_profiles(samples, document).model_dump(
            mode="json", exclude_none=True
        ),
        "rules": _validated_rule_dicts(rules),
    }


def _build_render_profiles(
    samples: dict[str, SampledParagraph],
    document: Document | None,
) -> RenderProfiles:
    fallback = RenderProfiles.fallback()
    front_profiles = dict(fallback.front_matter)

    role_map = {
        "abstract_body": SemanticRole.ABSTRACT_BODY,
        "keywords": SemanticRole.KEYWORDS,
        "abstract_body_en": SemanticRole.ABSTRACT_BODY_EN,
        "keywords_en": SemanticRole.KEYWORDS_EN,
    }
    for sample_role, semantic_role in role_map.items():
        sample = samples.get(sample_role)
        if sample is None:
            continue
        front_profiles[semantic_role] = _front_profile_from_sample(
            semantic_role,
            sample,
            fallback.front_matter[semantic_role],
        )

    return RenderProfiles(
        toc_visual=_toc_profile_from_document(document),
        front_matter=front_profiles,
        reference=ReferenceProfile(),
        acknowledgment=AcknowledgmentProfile(),
    )


def _front_profile_from_sample(
    semantic_role: SemanticRole,
    sample: SampledParagraph,
    fallback: FrontMatterProfile,
) -> FrontMatterProfile:
    first_run = next((run for run in sample.paragraph.runs if run.text.strip()), None)
    prefix_label = fallback.prefix_label
    if first_run is not None:
        prefix_label = first_run.text.strip() or prefix_label
        if semantic_role == SemanticRole.ABSTRACT_BODY and not prefix_label.startswith("摘要"):
            prefix_label = "摘要"
        elif semantic_role == SemanticRole.KEYWORDS and not prefix_label.startswith("关键词"):
            prefix_label = "关键词"
        elif semantic_role == SemanticRole.ABSTRACT_BODY_EN:
            prefix_label = "ABSTRACT"
        elif semantic_role == SemanticRole.KEYWORDS_EN:
            prefix_label = "KEY WORDS"

    return FrontMatterProfile(
        semantic_role=semantic_role,
        prefix_label=prefix_label,
        label_font_bold=bool(first_run.bold) if first_run is not None and first_run.bold is not None else fallback.label_font_bold,
        label_font_name=first_run.font.name if first_run is not None and first_run.font.name else fallback.label_font_name,
        label_font_size_pt=first_run.font.size.pt if first_run is not None and first_run.font.size is not None else fallback.label_font_size_pt,
        content_style=fallback.content_style,
        no_first_line_indent=True,
    )


def _toc_profile_from_document(document: Document | None) -> TocVisualProfile:
    fallback = TocVisualProfile()
    if document is None:
        return fallback

    style_names = {1: "TOC 1", 2: "TOC 2", 3: "TOC 3"}
    right_tab_twips: int | None = None
    has_dot_leader = True
    for level, style_name in style_names.items():
        if style_name not in document.styles:
            continue
        style = document.styles[style_name]
        tabs = style.element.findall(".//" + qn("w:tab"))
        for tab in tabs:
            if tab.get(qn("w:val")) == "right":
                right_tab_twips = int(tab.get(qn("w:pos")) or 0) or right_tab_twips
                has_dot_leader = tab.get(qn("w:leader")) == "dot"
                break
    return TocVisualProfile(
        has_dot_leader=has_dot_leader,
        page_number_alignment="right",
        right_tab_twips=right_tab_twips,
        entry_styles=style_names,
    )


def _missing_front_matter_default_rules(
    existing_rules: list[dict[str, Any]],
) -> list[dict[str, Any]]:
    existing_ids = {rule["id"] for rule in existing_rules}
    return [
        rule
        for rule in _front_matter_default_rules()
        if rule["id"] not in existing_ids
    ]


def _front_matter_default_rules() -> list[dict[str, Any]]:
    body_like = {
        "font_name": "Times New Roman",
        "east_asia_font": "宋体",
        "ascii_font": "Times New Roman",
        "hansi_font": "Times New Roman",
        "font_size_pt": 12.0,
        "alignment": "justify",
        "space_after_pt": 0.0,
        "first_line_indent_pt": 24.0,
        "line_spacing_pt": 23.0,
    }
    front_like = {**body_like, "first_line_indent_pt": 0.0}
    return [
        {
            "id": "title_en",
            "selector": {"semantic_role": "title_en"},
            "apply": {
                **front_like,
                "style_name": "FlowTitleEN",
                "alignment": "center",
                "font_size_pt": 18.0,
                "bold": True,
            },
            "priority": 20,
        },
        {
            "id": "abstract_body",
            "selector": {"semantic_role": "abstract_body"},
            "apply": {**front_like, "style_name": "FlowAbstractBody"},
            "priority": 20,
        },
        {
            "id": "keywords",
            "selector": {"semantic_role": "keywords"},
            "apply": {
                **body_like,
                "style_name": "FlowKeywords",
                "first_line_indent_pt": 0.0,
            },
            "priority": 20,
        },
        {
            "id": "abstract_body_en",
            "selector": {"semantic_role": "abstract_body_en"},
            "apply": {**front_like, "style_name": "FlowAbstractBodyEN"},
            "priority": 20,
        },
        {
            "id": "keywords_en",
            "selector": {"semantic_role": "keywords_en"},
            "apply": {
                **body_like,
                "style_name": "FlowKeywordsEN",
                "first_line_indent_pt": 0.0,
            },
            "priority": 20,
        },
        {
            "id": "references_item",
            "selector": {"semantic_role": "references_item"},
            "apply": {
                "font_name": "宋体",
                "east_asia_font": "宋体",
                "ascii_font": "Times New Roman",
                "hansi_font": "Times New Roman",
                "font_size_pt": 12.0,
                "right_indent_pt": 24.0,
                "line_spacing_pt": 23.0,
                "style_name": "FlowReferencesItem",
            },
            "priority": 20,
        },
        {
            "id": "acknowledgment_body",
            "selector": {"semantic_role": "acknowledgment_body"},
            "apply": {
                "font_name": "宋体",
                "east_asia_font": "宋体",
                "ascii_font": "Times New Roman",
                "hansi_font": "Times New Roman",
                "font_size_pt": 12.0,
                "alignment": "justify",
                "first_line_indent_pt": 24.0,
                "right_indent_pt": 24.0,
                "line_spacing_pt": 23.0,
                "style_name": "FlowAcknowledgmentBody",
            },
            "priority": 20,
        },
    ]


def _role_specs() -> dict[str, dict[str, Any]]:
    return {
        "title_cn": {
            "id": "title_cn",
            "label": "中文主标题",
            "description": "论文题名页中的中文主标题。",
            "selector": {"semantic_role": "title_cn"},
            "priority": 30,
            "style_name": "FlowTitleCN",
            "confidence": 0.78,
        },
        "h1": {
            "id": "heading_level_1",
            "label": "一级标题",
            "description": "正文一级章节标题。",
            "selector": {"node_kind": "heading", "level": 1},
            "priority": 1,
            "style_name": "FlowHeading1",
            "outline_level": 0,
            "confidence": 0.72,
        },
        "h2": {
            "id": "heading_level_2",
            "label": "二级标题",
            "description": "正文二级小节标题。",
            "selector": {"node_kind": "heading", "level": 2},
            "priority": 1,
            "style_name": "FlowHeading2",
            "outline_level": 1,
            "confidence": 0.68,
        },
        "body": {
            "id": "body_paragraph",
            "label": "正文段落",
            "description": "正文区域中的普通段落。",
            "selector": {"node_kind": "paragraph", "region": "body"},
            "priority": 1,
            "style_name": "FlowBody",
            "confidence": 0.72,
        },
        "title_en": {
            "id": "title_en",
            "label": "英文标题",
            "description": "论文题名页中的英文主标题。",
            "selector": {"semantic_role": "title_en"},
            "priority": 25,
            "style_name": "FlowTitleEN",
            "confidence": 0.68,
        },
        "abstract_body": {
            "id": "abstract_body",
            "label": "中文摘要正文",
            "description": "中文摘要段落。",
            "selector": {"semantic_role": "abstract_body"},
            "priority": 20,
            "style_name": "FlowAbstractBody",
            "confidence": 0.66,
        },
        "keywords": {
            "id": "keywords",
            "label": "中文关键词",
            "description": "中文关键词段落。",
            "selector": {"semantic_role": "keywords"},
            "priority": 20,
            "style_name": "FlowKeywords",
            "confidence": 0.66,
        },
        "abstract_body_en": {
            "id": "abstract_body_en",
            "label": "英文摘要正文",
            "description": "英文摘要段落。",
            "selector": {"semantic_role": "abstract_body_en"},
            "priority": 20,
            "style_name": "FlowAbstractBodyEN",
            "confidence": 0.66,
        },
        "keywords_en": {
            "id": "keywords_en",
            "label": "英文关键词",
            "description": "英文关键词段落。",
            "selector": {"semantic_role": "keywords_en"},
            "priority": 20,
            "style_name": "FlowKeywordsEN",
            "confidence": 0.66,
        },
        "references_heading": {
            "id": "references_heading",
            "label": "参考文献标题",
            "description": "参考文献区域的标题。",
            "selector": {"semantic_role": "references"},
            "priority": 10,
            "style_name": "FlowReferencesHeading",
            "outline_level": 0,
            "confidence": 0.9,
        },
        "back_paragraph": {
            "id": "back_paragraph",
            "label": "参考文献条目",
            "description": "参考文献区域中的条目正文。",
            "selector": {"node_kind": "paragraph", "region": "back"},
            "priority": 1,
            "style_name": "FlowBackParagraph",
            "confidence": 0.86,
        },
    }


def _front_matter_sample_role(text: str, title_cn_found: bool) -> str | None:
    compact = re.sub(r"\s+", "", text).lower()
    if compact.startswith("摘要") and len(compact) > 2:
        return "abstract_body"
    if compact.startswith("关键词"):
        return "keywords"
    if compact.startswith("abstract") and len(compact) > 8:
        return "abstract_body_en"
    if compact.startswith("keywords"):
        return "keywords_en"
    if title_cn_found and _looks_like_title_en(text):
        return "title_en"
    return None


def _looks_like_title_en(text: str) -> bool:
    if not (12 <= len(text) <= 180):
        return False
    if not re.search(r"[A-Za-z]", text):
        return False
    if re.search(r"[\u4e00-\u9fff]", text):
        return False
    if text.lower().startswith(("abstract", "keywords", "key words")):
        return False
    return not text.endswith((".", "!", "?"))


def extract_template_to_yaml(template_path: Path, output_path: Path) -> dict[str, SampledParagraph]:

    from flow_ai_open.templates.format_extractor import (
        catalog_to_profile_entries,
        catalog_to_rule_definitions,
        extract_format_catalog,
    )

    catalog = extract_format_catalog(template_path)
    profile_rules = catalog_to_profile_entries(catalog)
    compile_rules = catalog_to_rule_definitions(catalog)

    document = Document(template_path)
    legacy_samples = sniff_semantic_samples(document)
    legacy_profile = build_profile_document(legacy_samples, WordStyleResolver(document), document)

    legacy_profile["profile"]["schema_version"] = "format_catalog.v1"
    legacy_profile["profile"]["title"] = "Extracted format catalog"
    legacy_profile["profile"]["description"] = (
        "Phase 2 视觉指纹聚类 + 语义锚点对齐产出的格式菜单草案。"
    )
    legacy_profile["profile"]["rules"] = profile_rules
    legacy_profile["profile"]["format_catalog"] = catalog.model_dump(mode="json")
    legacy_profile["profile"]["warnings"] = _catalog_warnings(catalog)
    legacy_profile["rules"] = _validated_rule_dicts(compile_rules + _generated_support_rules())

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as file:
        yaml.safe_dump(
            legacy_profile,
            file,
            allow_unicode=True,
            sort_keys=False,
            default_flow_style=False,
        )
    return legacy_samples


def _catalog_warnings(catalog) -> list[str]:
    warnings: list[str] = []
    for slot in catalog.slots:
        if slot.confidence < 0.3:
            warnings.append(f"低置信度格式项: {slot.label} ({slot.slot_id})")
        if slot.source in {"missing", "opaque"}:
            warnings.append(f"缺少段落样本: {slot.label} ({slot.slot_id})")
    if catalog.opaque_slots:
        warnings.append(f"检测到 opaque/域块: {', '.join(catalog.opaque_slots)}")
    return warnings


def extract_template_to_yaml_legacy(template_path: Path, output_path: Path) -> dict[str, SampledParagraph]:

    document = Document(template_path)
    resolver = WordStyleResolver(document)
    samples = sniff_semantic_samples(document)
    profile_document = build_profile_document(samples, resolver, document)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    with output_path.open("w", encoding="utf-8") as file:
        yaml.safe_dump(
            profile_document,
            file,
            allow_unicode=True,
            sort_keys=False,
            default_flow_style=False,
        )
    return samples


def _find_title_cn(
    paragraphs: list[tuple[int, Paragraph, str]],
    search_limit: int = 40,
) -> SampledParagraph | None:
    candidates: list[tuple[float, int, Paragraph, str]] = []
    for index, paragraph, text in paragraphs[:search_limit]:
        if EXCLUDED_TITLE_RE.match(text):
            continue
        if len(text) < 6 or len(text) > 80:
            continue
        if not re.search(r"[\u4e00-\u9fff]", text):
            continue
        props = extract_physical_properties(paragraph)
        score = 0.0
        if props.get("alignment") == "center":
            score += 5
        if props.get("bold") is True:
            score += 2
        score += float(props.get("font_size_pt") or 0)
        score -= index * 0.05
        candidates.append((score, index, paragraph, text))

    if not candidates:
        return None

    _, index, paragraph, text = max(candidates, key=lambda item: item[0])
    return SampledParagraph("title_cn", paragraph, index, text)


def _looks_like_h1(paragraph: Paragraph, text: str) -> bool:
    if REFERENCE_RE.match(text) or EXCLUDED_TITLE_RE.match(text):
        return False
    outline_level = _extract_outline_level(paragraph)
    return outline_level == 0 or bool(H1_RE.match(text))


def _looks_like_h2(paragraph: Paragraph, text: str) -> bool:
    if REFERENCE_RE.match(text) or EXCLUDED_TITLE_RE.match(text):
        return False
    outline_level = _extract_outline_level(paragraph)
    return outline_level == 1 or bool(H2_RE.match(text))


def _looks_like_body(paragraph: Paragraph, text: str) -> bool:
    if len(text) < 80:
        return False
    if REFERENCE_RE.match(text) or "摘要" in text[:20] or "关键词" in text[:20]:
        return False
    if _looks_like_h1(paragraph, text) or _looks_like_h2(paragraph, text):
        return False
    if _looks_like_reference_item(text):
        return False
    if EXCLUDED_TITLE_RE.match(text):
        return False
    return bool(re.search(r"[\u4e00-\u9fff]", text))


def _find_body_sample(
    paragraphs: list[tuple[int, Paragraph, str]],
    h1_index: int | None,
    references_index: int | None,
) -> SampledParagraph | None:

    candidates: list[tuple[float, int, Paragraph, str]] = []
    start = h1_index if h1_index is not None else 0
    stop = references_index if references_index is not None else 10**9

    for index, paragraph, text in paragraphs:
        if index <= start or index >= stop:
            continue
        if not _looks_like_body(paragraph, text):
            continue

        props = extract_physical_properties(paragraph)
        score = min(len(text), 500) / 100
        if props.get("alignment") == "justify":
            score += 2
        if props.get("first_line_indent_pt") is not None:
            score += 2
        if props.get("bold") is True:
            score -= 3
        if props.get("font_size_pt") and float(props["font_size_pt"]) > 13:
            score -= 2
        candidates.append((score, index, paragraph, text))

    if not candidates:
        return None
    _, index, paragraph, text = max(candidates, key=lambda item: item[0])
    return SampledParagraph("body", paragraph, index, text)


def _find_heading_sample(
    paragraphs: list[tuple[int, Paragraph, str]],
    h1_index: int | None,
    references_index: int | None,
    level: int,
) -> SampledParagraph | None:

    start = h1_index if h1_index is not None else 0
    stop = references_index if references_index is not None else 10**9
    for index, paragraph, text in paragraphs:
        if index <= start or index >= stop:
            continue
        if level == 2 and _looks_like_h2(paragraph, text):
            return SampledParagraph("h2", paragraph, index, text)
    return None


def _looks_like_reference_item(text: str) -> bool:
    return bool(REFERENCE_ITEM_RE.match(text)) and len(text) >= 20


def _extract_font_properties(paragraph: Paragraph) -> dict[str, Any]:
    runs = [run for run in paragraph.runs if run.text.strip()]
    style_font = paragraph.style.font if paragraph.style is not None else None
    font_names: Counter[str] = Counter()
    east_asia_fonts: Counter[str] = Counter()
    ascii_fonts: Counter[str] = Counter()
    hansi_fonts: Counter[str] = Counter()
    sizes: Counter[float] = Counter()
    bold_votes: list[bool] = []

    for run in runs:
        weight = max(len(run.text.strip()), 1)
        run_font = run.font
        if run_font.name:
            font_names[run_font.name] += weight
        if run_font.size is not None:
            sizes[round(float(run_font.size.pt), 2)] += weight
        if run_font.bold is not None:
            bold_votes.extend([bool(run_font.bold)] * weight)

        r_fonts = run._element.rPr.rFonts if run._element.rPr is not None else None
        if r_fonts is not None:
            _count_xml_font(r_fonts, "w:eastAsia", east_asia_fonts, weight)
            _count_xml_font(r_fonts, "w:ascii", ascii_fonts, weight)
            _count_xml_font(r_fonts, "w:hAnsi", hansi_fonts, weight)

    properties: dict[str, Any] = {}
    if not font_names and style_font is not None and style_font.name:
        font_names[style_font.name] += 1
    if not sizes and style_font is not None and style_font.size is not None:
        sizes[round(float(style_font.size.pt), 2)] += 1
    if not bold_votes and style_font is not None and style_font.bold is not None:
        bold_votes.append(bool(style_font.bold))

    _put_counter_mode(properties, "font_name", font_names)
    _put_counter_mode(properties, "east_asia_font", east_asia_fonts)
    _put_counter_mode(properties, "ascii_font", ascii_fonts)
    _put_counter_mode(properties, "hansi_font", hansi_fonts)
    _put_counter_mode(properties, "font_size_pt", sizes)
    if bold_votes:
        properties["bold"] = sum(bold_votes) >= len(bold_votes) / 2

    if "east_asia_font" not in properties and "font_name" in properties:
        properties["east_asia_font"] = properties["font_name"]
    if "ascii_font" not in properties and "font_name" in properties:
        properties["ascii_font"] = properties["font_name"]
    if "hansi_font" not in properties and "font_name" in properties:
        properties["hansi_font"] = properties["font_name"]

    return properties


def _count_xml_font(counter_source: Any, attr_name: str, counter: Counter[str], weight: int) -> None:
    value = counter_source.get(qn(attr_name))
    if value:
        counter[value] += weight


def _put_counter_mode(properties: dict[str, Any], key: str, counter: Counter[Any]) -> None:
    if counter:
        properties[key] = counter.most_common(1)[0][0]


def _put_pt(properties: dict[str, Any], key: str, value: Any) -> None:
    if value is not None:
        properties[key] = round(float(value.pt), 2)


def _put_line_spacing(properties: dict[str, Any], value: Any) -> None:
    if value is None:
        return
    if isinstance(value, float):
        properties["line_spacing_multiple"] = round(float(value), 3)
        return
    if isinstance(value, Length):
        properties["line_spacing_pt"] = round(float(value.pt), 2)


def _extract_outline_level(paragraph: Paragraph) -> int | None:
    paragraph_properties = paragraph._p.pPr
    style_properties = (
        paragraph.style._element.pPr if paragraph.style is not None else None
    )

    for properties in (paragraph_properties, style_properties):
        if properties is None:
            continue
        outline = properties.find(qn("w:outlineLvl"))
        if outline is not None:
            value = outline.get(qn("w:val"))
            if value is not None:
                return int(value)
    return None


def _drop_none(properties: dict[str, Any]) -> dict[str, Any]:
    return {key: value for key, value in properties.items() if value is not None}


def _validated_rule_dicts(rules: list[dict[str, Any]]) -> list[dict[str, Any]]:
    return [
        RuleNode.model_validate(rule).model_dump(mode="json", exclude_none=True)
        for rule in rules
    ]


def _intent_to_readable(intent: dict[str, Any]) -> list[str]:
    readable: list[str] = []
    base_font_size = float(intent.get("font_size_pt") or 12.0)
    if font := intent.get("east_asia_font") or intent.get("font_name"):
        readable.append(f"中文字体：{font}")
    if font := intent.get("ascii_font") or intent.get("hansi_font"):
        readable.append(f"西文字体：{font}")
    if size := intent.get("font_size_pt"):
        readable.append(f"字号：{pt_to_cn_size(float(size))}")
    if intent.get("bold") is True:
        readable.append("字形：加粗")
    elif intent.get("bold") is False:
        readable.append("字形：不加粗")
    if alignment := intent.get("alignment"):
        readable.append(f"对齐：{_alignment_label(alignment)}")
    if value := intent.get("first_line_indent_pt"):
        readable.append(
            f"首行缩进：{pt_to_cn_chars(float(value), base_font_size_pt=base_font_size)}"
        )
    if value := intent.get("hanging_indent_pt"):
        readable.append(
            f"悬挂缩进：{pt_to_cn_chars(float(value), base_font_size_pt=base_font_size)}"
        )
    if value := intent.get("left_indent_pt"):
        readable.append(
            f"左缩进：{pt_to_cn_chars(float(value), base_font_size_pt=base_font_size)}"
        )
    if value := intent.get("right_indent_pt"):
        readable.append(
            f"右缩进：{pt_to_cn_chars(float(value), base_font_size_pt=base_font_size)}"
        )
    if value := intent.get("space_before_pt"):
        readable.append(f"段前：{float(value):.1f} 磅")
    if value := intent.get("space_after_pt"):
        readable.append(f"段后：{float(value):.1f} 磅")
    if value := intent.get("line_spacing_multiple"):
        readable.append(f"行距：{line_spacing_to_cn(multiple=float(value))}")
    if value := intent.get("line_spacing_pt"):
        readable.append(f"行距：{line_spacing_to_cn(fixed_pt=float(value))}")
    if intent.get("page_break_before") is True:
        readable.append("分页：段前分页")
    if intent.get("keep_with_next") is True:
        readable.append("分页：与下段同页")
    if value := intent.get("outline_level"):
        readable.append(f"大纲级别：{int(value) + 1} 级")
    elif intent.get("outline_level") == 0:
        readable.append("大纲级别：1 级")
    if style_name := intent.get("style_name"):
        readable.append(f"内部样式：{style_name}")
    return readable


def _alignment_label(value: str) -> str:
    return {
        "left": "左对齐",
        "center": "居中",
        "right": "右对齐",
        "justify": "两端对齐",
        "distribute": "分散对齐",
        "unknown": "未知",
    }.get(value, value)


def _profile_warnings(entries: list[RuleProfileEntry]) -> list[str]:
    warnings: list[str] = []
    by_id = {entry.id: entry for entry in entries}
    for required in ("title_cn", "heading_level_1", "body_paragraph"):
        if required not in by_id:
            warnings.append(f"未提取到关键规则：{required}")
    for entry in entries:
        if not any(
            key in entry.intent
            for key in ("font_name", "east_asia_font", "ascii_font", "font_size_pt")
        ):
            warnings.append(f"{entry.label} 缺少字体/字号信息，建议人工确权。")
    return warnings


def _generated_support_rules() -> list[dict[str, Any]]:
    return [
        {
            "id": "generated_toc_anchor",
            "selector": {"node_kind": "generated_anchor", "anchor_type": "toc"},
            "apply": {
                "generated_anchor_type": "toc",
                "generated_anchor_text": "[TOC will be generated here]",
                "style_name": "FlowTocPlaceholder",
                "font_name": "SimHei",
                "east_asia_font": "SimHei",
                "ascii_font": "Times New Roman",
                "hansi_font": "Times New Roman",
                "font_size_pt": 12,
                "bold": True,
                "alignment": "center",
                "space_before_pt": 12,
                "space_after_pt": 12,
            },
            "priority": 20,
        },
        {
            "id": "opaque_placeholder",
            "selector": {"node_kind": "opaque"},
            "apply": {
                "opaque_placeholder_text": "[UNSUPPORTED OPAQUE: {opaque_type}]",
                "style_name": "FlowOpaquePlaceholder",
                "font_name": "SimSun",
                "east_asia_font": "SimSun",
                "ascii_font": "Times New Roman",
                "hansi_font": "Times New Roman",
                "font_size_pt": 10.5,
                "alignment": "center",
                "space_before_pt": 6,
                "space_after_pt": 6,
            },
            "priority": 1,
        },
    ]


if __name__ == "__main__":
    root = Path(__file__).resolve().parents[2]
    template = root / "template_samples.docx"
    output = root / "src" / "templates" / "extracted_thesis.yaml"

    sniffed_samples = extract_template_to_yaml(template, output)
    print("=== Template Sniffing Samples ===")
    for role in ("title_cn", "h1", "h2", "body", "references_heading", "back_paragraph"):
        sample = sniffed_samples.get(role)
        if sample is None:
            print(f"[MISS][{role}]")
            continue
        preview = sample.text[:120]
        print(f"[HIT][{role}][#{sample.index}] {preview}")
    profile_data = yaml.safe_load(output.read_text(encoding="utf-8"))
    print("=== Rule Profile (Plain Language) ===")
    for rule in profile_data.get("profile", {}).get("rules", []):
        print(f"- {rule['label']} ({rule['id']})")
        for line in rule.get("readable", []):
            print(f"  · {line}")
        sample = rule.get("sample", {})
        print(f"  · 样本：#{sample.get('index')} {sample.get('text', '')[:80]}")
    warnings = profile_data.get("profile", {}).get("warnings", [])
    if warnings:
        print("=== Profile Warnings ===")
        for warning in warnings:
            print(f"! {warning}")
    print(f"Saved YAML: {output}")
